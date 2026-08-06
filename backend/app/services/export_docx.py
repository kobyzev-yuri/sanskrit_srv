"""Build a downloadable DOCX from project page HTML.

Modes: text (HTML only) | interleave (scan image then HTML for each source page).
"""
from __future__ import annotations

import io
import re
import uuid
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from app.services.storage import ensure_dirs

_API_IMG_RE = re.compile(
    r"/api/v1/pages/([0-9a-fA-F-]{36})/figures/([^\"'\s>]+)",
    re.I,
)

_FONT = "Noto Serif Devanagari"
_FALLBACK_FONT = "FreeSerif"
_SCAN_MAX_PX = 1400
_SCAN_JPEG_Q = 78
# Match compact PDF measure so one source page ≈ one Word page.
_BODY_PT = 8.0
_H1_PT = 10.0
_SMALL_PT = 7.0


def build_project_docx(
    project_id: uuid.UUID,
    slug: str,
    title: str,
    pages: list[tuple[int, str, str | None]],
    *,
    title_sa: str | None = None,
    mode: str = "text",
) -> Path:
    """pages: list of (page_no, html_fragment, scan_path|None).

    mode=text — HTML only.
    mode=interleave — scan plate, then HTML, for each source page.
    """
    mode = (mode or "text").strip().lower()
    if mode not in ("text", "interleave"):
        mode = "text"

    out_dir = ensure_dirs() / "exports" / str(project_id)
    out_dir.mkdir(parents=True, exist_ok=True)
    suffix = "-interleave" if mode == "interleave" else ""
    out_path = out_dir / f"{slug}{suffix}.docx"

    doc = Document()
    _setup_styles(doc)
    _add_cover(doc, title, title_sa)  # ends with a page break

    body_n = 0
    need_break = False  # cover already broke; first body sheet follows cleanly
    for page_no, html, scan_path in pages:
        frag = (html or "").strip()
        blocks: list[tuple[str, object]] = []
        if mode == "interleave" and scan_path and Path(scan_path).is_file():
            blocks.append(("scan", Path(scan_path)))
        if frag:
            blocks.append(("html", frag))
        elif mode == "text":
            continue
        if not blocks:
            continue

        for kind, payload in blocks:
            if need_break:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            need_break = True
            if kind == "scan":
                _add_scan_image(doc, payload)  # type: ignore[arg-type]
            else:
                _html_to_docx(doc, str(payload), project_id, page_no)
            body_n += 1

    if body_n == 0:
        p = doc.add_paragraph()
        _set_run(p.add_run("(нет страниц с текстом для выгрузки)"), size=_BODY_PT)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tmp_out = out_path.with_suffix(f".{uuid.uuid4().hex}.tmp.docx")
    doc.save(tmp_out.as_posix())
    tmp_out.replace(out_path)
    return out_path


def _setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    # ~412×612 pt Indian crown (taller than A5)
    section.page_width = Cm(14.55)
    section.page_height = Cm(21.6)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = Pt(_BODY_PT)
    normal.font.color.rgb = RGBColor(0x1A, 0x18, 0x14)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.28
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.space_before = Pt(0)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), _FONT)
    rfonts.set(qn("w:hAnsi"), _FONT)
    rfonts.set(qn("w:cs"), _FONT)
    rfonts.set(qn("w:eastAsia"), _FONT)


def _add_cover(doc: Document, title: str, title_sa: str | None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p.add_run(title or "Untitled"), size=13, bold=True)
    if title_sa and title_sa.strip():
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run(p2.add_run(title_sa.strip()), size=10.5)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = _set_run(p3.add_run("Sanskrit SRV"), size=_SMALL_PT)
    r.font.color.rgb = RGBColor(0x6B, 0x65, 0x60)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _add_scan_image(doc: Document, scan_path: Path) -> None:
    jpeg = _scan_jpeg_bytes(scan_path)
    if not jpeg:
        try:
            doc.add_picture(scan_path.as_posix(), width=Inches(5.0))
        except Exception:  # noqa: BLE001
            p = doc.add_paragraph()
            _set_run(p.add_run(f"[scan missing: {scan_path.name}]"), size=_SMALL_PT)
        return
    stream = io.BytesIO(jpeg)
    # ~ usable width inside A5 margins
    doc.add_picture(stream, width=Inches(5.0))


def _scan_jpeg_bytes(scan_path: Path) -> bytes | None:
    try:
        from PIL import Image

        with Image.open(scan_path) as img:
            img = img.convert("RGB")
            img.thumbnail((_SCAN_MAX_PX, _SCAN_MAX_PX), Image.Resampling.LANCZOS)
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=_SCAN_JPEG_Q, optimize=True)
            return buf.getvalue()
    except Exception:  # noqa: BLE001
        return None


def _html_to_docx(
    doc: Document, html: str, project_id: uuid.UUID, page_no: int
) -> None:
    parser = _HtmlToDocx(doc, project_id, page_no)
    parser.feed(html)
    parser.close()
    parser.flush()


def _set_run(run, *, size: float = _BODY_PT, bold: bool = False):
    run.font.name = _FONT
    run.font.size = Pt(size)
    run.bold = bold
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), _FONT)
    rfonts.set(qn("w:hAnsi"), _FONT)
    rfonts.set(qn("w:cs"), _FONT)
    rfonts.set(qn("w:eastAsia"), _FALLBACK_FONT)
    return run


def _resolve_figure_path(
    src: str, project_id: uuid.UUID, page_no: int
) -> Path | None:
    from app.services.layout_assets import figure_file

    if not src:
        return None
    m = _API_IMG_RE.search(src)
    name = m.group(2) if m else Path(src.split("?")[0]).name
    if re.fullmatch(r"(emb|crop)-\d{2}\.png", name or ""):
        return figure_file(project_id, page_no, name)
    return None


class _HtmlToDocx(HTMLParser):
    """Map our page HTML fragment into python-docx paragraphs / tables / images."""

    _BLOCK = frozenset(
        {"p", "h1", "h2", "h3", "footer", "li", "div", "section", "article", "figcaption"}
    )

    def __init__(self, doc: Document, project_id: uuid.UUID, page_no: int):
        super().__init__(convert_charrefs=True)
        self.doc = doc
        self.project_id = project_id
        self.page_no = page_no
        self._para = None
        self._align = WD_ALIGN_PARAGRAPH.LEFT
        self._size = _BODY_PT
        self._bold = False
        self._skip = 0
        self._table = None
        self._row = None
        self._cell_paras: list = []
        self._in_td = False
        self._classes: list[str] = []
        self._pending_text: list[str] = []

    def flush(self) -> None:
        self._end_para()

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        ad = {k.lower(): (v or "") for k, v in attrs}
        classes = (ad.get("class") or "").split()

        if tag in ("script", "style"):
            self._skip += 1
            return
        if self._skip:
            return

        if tag == "br":
            if self._para is not None:
                self._para.add_run().add_break()
            return

        if tag == "img":
            self._end_para()
            path = _resolve_figure_path(ad.get("src", ""), self.project_id, self.page_no)
            if path and path.is_file():
                try:
                    self.doc.add_picture(path.as_posix(), width=Inches(4.2))
                except Exception:  # noqa: BLE001
                    pass
            return

        if tag == "table":
            self._end_para()
            self._table = {"rows": []}
            return
        if tag == "tr" and self._table is not None:
            self._row = []
            self._table["rows"].append(self._row)
            return
        if tag in ("td", "th") and self._row is not None:
            self._in_td = True
            self._cell_paras = [""]
            self._bold = tag == "th"
            return

        if tag in self._BLOCK:
            self._end_para()
            self._classes = classes
            self._align = WD_ALIGN_PARAGRAPH.LEFT
            self._size = _BODY_PT
            self._bold = False
            if tag == "h1" or "type-lg" in classes:
                self._size = _H1_PT
                self._bold = True
            elif "type-sm" in classes or "running-head" in classes or "page-num" in classes:
                self._size = _SMALL_PT
            elif "type-md" in classes:
                self._size = _BODY_PT
            if any(
                c in classes
                for c in ("centered", "shloka", "running-head", "page-num", "cover")
            ) or tag in ("h1", "footer"):
                self._align = WD_ALIGN_PARAGRAPH.CENTER
            # Start paragraph on first text / leave for empty skip
            self._para = None
            return

        if tag in ("strong", "b"):
            self._bold = True
        if tag == "figure":
            self._end_para()

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in ("script", "style"):
            self._skip = max(0, self._skip - 1)
            return
        if self._skip:
            return

        if tag in ("td", "th") and self._in_td and self._row is not None:
            self._row.append("".join(self._cell_paras).strip())
            self._in_td = False
            self._bold = False
            return

        if tag == "table" and self._table is not None:
            self._emit_table(self._table["rows"])
            self._table = None
            self._row = None
            return

        if tag in self._BLOCK:
            self._end_para()
            self._classes = []
            return

        if tag in ("strong", "b"):
            self._bold = False

    def handle_data(self, data: str) -> None:
        if self._skip or not data:
            return
        # HTML collapses newlines/indent inside <p>; Word does not — normalize.
        data = re.sub(r"\s+", " ", data)
        if self._in_td:
            self._cell_paras.append(data)
            return
        # Skip whitespace-only between blocks
        if self._para is None:
            data = data.lstrip()
            if not data:
                return
            self._para = self.doc.add_paragraph()
            self._para.alignment = self._align
            self._para.paragraph_format.space_after = Pt(2)
            self._para.paragraph_format.space_before = Pt(0)
            self._para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            self._para.paragraph_format.line_spacing = 1.28
            if "indent" in self._classes:
                self._para.paragraph_format.first_line_indent = Cm(0.5)
        if not data:
            return
        run = self._para.add_run(data)
        _set_run(run, size=self._size, bold=self._bold)

    def _end_para(self) -> None:
        self._para = None

    def _emit_table(self, rows: list[list[str]]) -> None:
        rows = [r for r in rows if any(c.strip() for c in r)]
        if not rows:
            return
        cols = max(len(r) for r in rows)
        table = self.doc.add_table(rows=len(rows), cols=cols)
        table.style = "Table Grid"
        for ri, row in enumerate(rows):
            for ci in range(cols):
                text = row[ci] if ci < len(row) else ""
                cell = table.rows[ri].cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                _set_run(p.add_run(text), size=_SMALL_PT)
        self.doc.add_paragraph()
